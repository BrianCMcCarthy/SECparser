from flask import Flask, request, jsonify, send_from_directory
import yfinance as yf
import pandas as pd
import numpy as np
import os
import fitz  # PyMuPDF
import re
from typing import List, Dict
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import io
import base64
from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from functools import wraps

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

API_KEY = os.environ.get("API_KEY", "supersecretkey")

def require_api_key(view_function):
    @wraps(view_function)
    def decorated_function(*args, **kwargs):
        key = request.headers.get("X-API-KEY")
        if key and key == API_KEY:
            return view_function(*args, **kwargs)
        else:
            return jsonify({"error": "Unauthorized"}), 401
    return decorated_function

# === Analysis & Helper Functions ===
def extract_latest(series, fallback=None):
    try:
        value = series.dropna().iloc[0]
        return int(value) if float(value).is_integer() else round(float(value), 2)
    except:
        return fallback

app = Flask(__name__)
app.config["UPLOAD_FOLDER"] = "uploads"
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

# === Aliases for Label Matching ===
FINANCIAL_LABEL_ALIASES = {
    "Revenue": ["Total Revenue", "Revenues", "Net Sales", "Sales"],
    "Gross Profit": ["Gross Profit", "Gross Income"],
    "SG&A": [
        "Selling General Administrative",
        "Operating Expenses",
        "SG&A Expense",
        "Selling & Admin",
        "Selling and Admin",
        "Selling, general and administrative expenses"
    ],
    "Net Income": ["Net Income", "Net Earnings", "Net Profit", "Income Available to Common Stockholders"],
    "Cash": ["Cash", "Cash and Cash Equivalents", "Cash & Equivalents"],
    "Total Debt": ["Long Term Debt", "Total Debt", "Total Liabilities"],
    "Equity": ["Total Stockholder Equity", "Total Equity", "Shareholder Equity"],
    "Operating Cash Flow": ["Total Cash From Operating Activities", "Net Cash Provided by Operating Activities"],
    "CapEx": ["Capital Expenditures", "Purchase of Property and Equipment"],
    "Buybacks": ["Repurchase Of Stock", "Share Repurchase", "Treasury Stock Purchased"]
}

# === Analysis & Helper Functions ===
def extract_latest(series, fallback=None):
    try:
        value = series.dropna().iloc[0]
        return int(value) if float(value).is_integer() else round(float(value), 2)
    except:
        return fallback

def safe_extract(df, key):
    labels = FINANCIAL_LABEL_ALIASES.get(key, [key])
    normalized_index = {i.lower(): i for i in df.index}
    for label in labels:
        label_lower = label.lower()
        if label_lower in normalized_index:
            return extract_latest(df.loc[normalized_index[label_lower]])
    print(f"[WARN] Missing: {key} - Checked: {labels}")
    return None

def generate_docx_report(ticker, summary, parsed):
    document = Document()
    document.add_heading(f"Activist Report: {ticker.upper()}", 0)

    document.add_heading("1. Executive Summary", level=1)
    document.add_paragraph("This report evaluates financial performance, strategic direction, and governance quality of the target company.")

    document.add_heading("2. Financial Highlights", level=1)
    for key in summary:
        value_entry = summary[key]
        if isinstance(value_entry, dict):
            document.add_paragraph(f"{key}: {value_entry['value']} [{value_entry['source']}]")
        else:
            document.add_paragraph(f"{key}: {value_entry}")

    document.add_heading("3. Charts", level=1)
    document.add_paragraph("Charts visualize key financial metrics and reveal capital structure, margin trends, and investment efficiency.")

    # Placeholder chart rendering - assuming plots exist
    chart_keys = ["SG&A", "Net Income", "Long Term Debt", "Share Buybacks"]
    for chart_key in chart_keys:
        chart_path = os.path.join(UPLOAD_FOLDER, f"{ticker}_{chart_key.replace(' ', '_')}.png")
        if os.path.exists(chart_path):
            document.add_heading(chart_key, level=2)
            document.add_picture(chart_path, width=Inches(6))

    document.add_heading("4. Governance & Board Review", level=1)
    board_table = parsed.get("board_comp_table", [])
    if board_table:
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        for idx, title in enumerate(["Name", "Title", "Amount", "Type", "Line", "Year"]):
            hdr_cells[idx].text = title
        for entry in board_table:
            row = table.add_row().cells
            row[0].text = entry.get("Name", "")
            row[1].text = entry.get("Title", "")
            row[2].text = entry.get("Amount", "")
            row[3].text = entry.get("Type", "")
            row[4].text = entry.get("Line", "")
            row[5].text = entry.get("Year", "")
    else:
        document.add_paragraph("No board compensation findings in uploaded materials.")

    document.add_heading("5. Strategic Positioning Flags", level=1)
    if parsed.get("strategy_flags"):
        for item in parsed["strategy_flags"]:
            document.add_paragraph(item)
    else:
        document.add_paragraph("No strategic initiative disclosures found in uploads.")

    file_path = os.path.join(UPLOAD_FOLDER, f"{ticker}_activist_report.docx")
    document.save(file_path)
    return file_path

def calculate_trends(df, line_item):
    aliases = FINANCIAL_LABEL_ALIASES.get(line_item, [line_item])
    for label in aliases:
        if label in df.index:
            values = df.loc[label].dropna().astype(float)
            if len(values) < 2:
                return None
            cagr = ((values[0] / values[-1]) ** (1 / (len(values) - 1)) - 1) * 100
            return round(cagr, 2)
    return None

def label_source(value, source):
    return {"value": value, "source": source if value is not None else "Missing"}

def analyze_company(ticker):
    try:
        stock = yf.Ticker(ticker)
        info = stock.info if stock.info else {}
        fin = stock.financials if not stock.financials.empty else pd.DataFrame()
        bal = stock.balance_sheet if not stock.balance_sheet.empty else pd.DataFrame()
        cf = stock.cashflow if not stock.cashflow.empty else pd.DataFrame()
        qbal = stock.quarterly_balance_sheet if not stock.quarterly_balance_sheet.empty else pd.DataFrame()
        qcf = stock.quarterly_cashflow if not stock.quarterly_cashflow.empty else pd.DataFrame()
    except Exception as e:
        return {"error": f"Yahoo Finance failed for {ticker}: {str(e)}"}

    name = info.get("longName", ticker)

    def get_val(df, label, fallback_label=None):
        try:
            if label in df.index:
                return extract_latest(df.loc[label])
            if fallback_label and fallback_label in df.index:
                return extract_latest(df.loc[fallback_label])
        except Exception:
            return None
        return None

    summary = {
        "Company": name,
        "Ticker": ticker.upper(),
        "Revenue": label_source(get_val(fin, "Total Revenue"), "Yahoo Finance"),
        "Gross Profit": label_source(get_val(fin, "Gross Profit"), "Yahoo Finance"),
        "SG&A": label_source(get_val(fin, "Selling General Administrative", "Operating Expenses"), "Estimated"),
        "Net Income": label_source(get_val(fin, "Net Income"), "Yahoo Finance")
    }

    rev = summary["Revenue"]["value"]
    gp = summary["Gross Profit"]["value"]
    ni = summary["Net Income"]["value"]
    sga = summary["SG&A"]["value"]

    summary["Gross Margin (%)"] = label_source(round(gp / rev * 100, 2) if gp and rev else None, "Calculated")
    summary["Net Income Margin (%)"] = label_source(round(ni / rev * 100, 2) if ni and rev else None, "Calculated")
    summary["SG&A as % of Revenue"] = label_source(round(sga / rev * 100, 2) if sga and rev else None, "Calculated")

    cash = get_val(bal, "Cash") or get_val(qbal, "Cash")
    debt = get_val(bal, "Long Term Debt", "Total Debt") or get_val(qbal, "Long Term Debt", "Total Debt")
    equity = get_val(bal, "Total Stockholder Equity") or get_val(qbal, "Total Stockholder Equity")

    summary["Cash"] = label_source(cash, "Yahoo Finance")
    summary["Total Debt"] = label_source(debt, "Yahoo Finance")
    summary["Net Debt"] = label_source(debt - cash if debt and cash else None, "Calculated")
    summary["Debt-to-Equity Ratio"] = label_source(round(debt / equity, 2) if debt and equity else None, "Calculated")

    ocf = get_val(cf, "Total Cash From Operating Activities") or get_val(qcf, "Total Cash From Operating Activities")
    capex = get_val(cf, "Capital Expenditures") or get_val(qcf, "Capital Expenditures")
    buybacks = get_val(cf, "Repurchase Of Stock") or get_val(qcf, "Repurchase Of Stock")

    summary["Operating Cash Flow"] = label_source(ocf, "Estimated")
    summary["CapEx"] = label_source(capex, "Estimated")
    summary["Share Buybacks"] = label_source(buybacks, "Estimated")

    fcf = ocf + capex if ocf and capex else None
    summary["Free Cash Flow"] = label_source(fcf, "Calculated")
    summary["FCF Margin (%)"] = label_source(round(fcf / rev * 100, 2) if fcf and rev else None, "Calculated")

    summary["Revenue CAGR (%)"] = label_source(calculate_trends(fin, "Total Revenue"), "Calculated")
    summary["Net Income CAGR (%)"] = label_source(calculate_trends(fin, "Net Income"), "Calculated")
    summary["SG&A CAGR (%)"] = label_source(calculate_trends(fin, "Selling General Administrative"), "Calculated")

    try:
        market_cap = info.get("marketCap", None)
        if market_cap and fcf and fcf != 0:
            irr_table = []
            hold_period = 3
            for multiple in range(8, 13):
                exit_ev = multiple * fcf
                entry_ev = market_cap + (debt or 0) - (cash or 0)
                irr = ((exit_ev / entry_ev) ** (1 / hold_period) - 1) * 100 if entry_ev > 0 else None
                irr_table.append({"Exit EV/FCF": multiple, "IRR (%)": round(irr, 2) if irr else None})
            summary["IRR Table"] = irr_table
    except Exception as e:
        summary["IRR Table"] = f"Error calculating IRR: {str(e)}"

    return summary

def parse_uploaded_content():
    parsed_data = {
        "board_insights": [],
        "strategy_flags": [],
        "board_comp_table": []
    }
    try:
        for filename in os.listdir(app.config["UPLOAD_FOLDER"]):
            if not filename.lower().endswith(".pdf"):
                continue
            filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
            try:
                doc = fitz.open(filepath)
                text = "\n".join([page.get_text() for page in doc])
                doc.close()
            except Exception as pdf_error:
                parsed_data["board_insights"].append(f"Failed to parse {filename}: {pdf_error}")
                continue

            for line in text.split("\n"):
                if re.search(r"(?i)director compensation|total compensation|meeting fees", line):
                    parsed_data["board_insights"].append(line.strip())
                    amt = re.search(r"\$\s?\d{1,3}(?:,\d{3})*(?:\.\d{2})?", line)
                    year = re.search(r"\b(20\d{2})\b", line)
                    parsed_data["board_comp_table"].append({
                        "Name": "Unknown",
                        "Title": "Director",
                        "Amount": amt.group(0) if amt else "-",
                        "Type": "Unknown",
                        "Line": line.strip(),
                        "Year": year.group(0) if year else "N/A"
                    })
                if re.search(r"(?i)FLX Rewards|loyalty program|strategic initiative|omnichannel", line):
                    parsed_data["strategy_flags"].append(line.strip())
    except Exception as e:
        parsed_data["error"] = f"Parse error: {str(e)}"

    return parsed_data

@app.route("/upload-file", methods=["POST"])
@require_api_key
def upload_file():
    if "file" not in request.files:
        return jsonify({"error": "No file part in request"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No selected file"}), 400

    filepath = os.path.join(app.config["UPLOAD_FOLDER"], file.filename)
    file.save(filepath)

    try:
        doc = fitz.open(filepath)
        full_text = "\n".join([page.get_text() for page in doc])
        doc.close()
    except Exception as e:
        return jsonify({"error": f"Failed to process file: {str(e)}"}), 500

    keywords = [
        "Board of Directors", "Compensation Committee", "Shareholder", "Dividend",
        "BOPIS", "Loyalty", "FLX Rewards", "Private Label", "Digital", "App", "Buyback",
        "Ometria", "CDP", "Return Policy", "Omnichannel", "E-commerce"
    ]

    excerpts = []
    for line in full_text.split("\n"):
        for kw in keywords:
            if kw.lower() in line.lower():
                excerpts.append({"keyword": kw, "excerpt": line.strip()})

    board_comp_table = extract_board_comp_table(full_text)

    return jsonify({
        "filename": file.filename,
        "excerpts": excerpts,
        "board_comp_table": board_comp_table,
        "keywords_matched": list(set([e["keyword"] for e in excerpts])),
        "num_findings": len(excerpts),
        "num_comp_entries": len(board_comp_table)
    })

def extract_board_comp_table(text: str) -> List[Dict[str, str]]:
    comp_entries = []
    pattern = r"(?i)(?:[\$€¥£]\s?[\d{1,3},]*\d{1,3}(?:\.\d{1,2})?)"
    lines = text.split("\n")
    for line in lines:
        matches = re.findall(pattern, line)
        for match in matches:
            normalized = match.replace(",", "").replace(" ", "")
            try:
                value = float(re.sub(r"[^\d.]", "", normalized))
                if value >= 1 and value <= 20000000:
                    comp_entries.append({
                        "Line": line.strip(),
                        "Reported Comp": f"${int(value) if value.is_integer() else round(value, 2)}"
                    })
            except:
                continue
    return comp_entries

@app.route("/generate-charts", methods=["GET"])
@require_api_key
def generate_charts():
    ticker = request.args.get("ticker")
    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    stock = yf.Ticker(ticker)
    fin = stock.financials
    cf = stock.cashflow
    bal = stock.balance_sheet

    charts = {}

    chart_targets = {
        "SG&A": ["Selling General Administrative", "Operating Expenses"],
        "Net Income": ["Net Income"],
        "Long Term Debt": ["Long Term Debt"],
        "Share Buybacks": ["Repurchase Of Stock"],
        "CapEx": ["Capital Expenditures"],
        "Operating Cash Flow": ["Total Cash From Operating Activities"],
        "Revenue": ["Total Revenue"]
    }

def plot_and_encode(series, title, ticker):
    series = series.dropna().astype(float)
    if series.empty or len(series) < 2:
        return None

    fig, ax = plt.subplots()
    series[::-1].plot(kind="bar", ax=ax, color="steelblue")
    ax.set_title(title, fontsize=14, fontweight='bold')
    ax.set_ylabel("USD", fontsize=12)
    ax.set_xlabel("Date", fontsize=12)
    ax.grid(True, which='major', axis='y', linestyle='--', alpha=0.7)
    ax.legend([title], loc='upper left', fontsize=10)
    for i, v in enumerate(series[::-1]):
        ax.text(i, v, f"{v:,.0f}", ha='center', va='bottom', fontsize=8, rotation=0)
    plt.xticks(rotation=45, ha='right')
    plt.tight_layout()

    chart_path = f"{UPLOAD_FOLDER}/{ticker}_{title.replace(' ', '_')}.png"
    plt.savefig(chart_path, format="png")
    plt.close(fig)
    return chart_path

    for label, options in chart_targets.items():
        for key in options:
            for df in [fin, cf, bal]:
                if key in df.index:
                    encoded = plot_and_encode(df.loc[key], label)
                    if encoded:
                        charts[label] = encoded
                        break
            if label in charts:
                break

    return jsonify(charts)

@app.route("/generate-docx", methods=["GET"])
@require_api_key
def generate_docx():
    ticker = request.args.get("ticker")
    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    summary = analyze_company(ticker.upper())
    parsed = parse_uploaded_content()

    file_path = generate_docx_report(ticker, summary, parsed)
    return jsonify({"doc_path": file_path})

@app.route("/generate-brief", methods=["GET"])
@require_api_key
def generate_brief():
    ticker = request.args.get("ticker")
    peers = request.args.get("peers", "")
    peer_list = [p.strip().upper() for p in peers.split(",") if p.strip()]

    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    main_summary = analyze_company(ticker.upper())
    peer_summaries = [analyze_company(p) for p in peer_list]
    parsed = parse_uploaded_content()

    insights = []
    main_rev = main_summary["Revenue"]["value"] or 1
    for peer in peer_summaries:
        if peer["Gross Margin (%)"]["value"] and main_summary["Gross Margin (%)"]["value"] and \
           peer["Gross Margin (%)"]["value"] > main_summary["Gross Margin (%)"]["value"] + 2:
            insights.append(f"{main_summary['Ticker']} gross margin ({main_summary['Gross Margin (%)']['value']}%) is below {peer['Ticker']} at {peer['Gross Margin (%)']['value']}%. [[SOURCE: {main_summary['Gross Margin (%)']['source']}]]")

        if peer["SG&A as % of Revenue"]["value"] and main_summary["SG&A as % of Revenue"]["value"] and \
           peer["SG&A as % of Revenue"]["value"] < main_summary["SG&A as % of Revenue"]["value"] - 2:
            insights.append(f"{main_summary['Ticker']} SG&A % of revenue ({main_summary['SG&A as % of Revenue']['value']}%) is higher than {peer['Ticker']} at {peer['SG&A as % of Revenue']['value']}%. [[SOURCE: {main_summary['SG&A as % of Revenue']['source']}]]")

        if peer["FCF Margin (%)"]["value"] and main_summary["FCF Margin (%)"]["value"] and \
           peer["FCF Margin (%)"]["value"] > main_summary["FCF Margin (%)"]["value"] + 2:
            insights.append(f"{main_summary['Ticker']} FCF margin ({main_summary['FCF Margin (%)']['value']}%) lags {peer['Ticker']} at {peer['FCF Margin (%)']['value']}%. [[SOURCE: {main_summary['FCF Margin (%)']['source']}]]")

    if parsed.get("strategy_flags"):
        insights.append("\nStrategic initiatives referenced in uploaded materials:")
        for line in parsed["strategy_flags"]:
            insights.append(f"- {line}")

    if parsed.get("board_insights"):
        insights.append("\nBoard governance references in uploaded materials:")
        for line in parsed["board_insights"]:
            insights.append(f"- {line}")

    return jsonify({
        "ticker": ticker,
        "executive_summary": insights,
        "parsed_files": parsed,
        "main_summary": main_summary,
        "peer_summaries": peer_summaries
    })

@app.route("/analyze-activist", methods=["GET"])
@require_api_key
def analyze_activist():
    ticker = request.args.get("ticker")
    peers = request.args.get("peers", "")
    peer_list = [p.strip().upper() for p in peers.split(",") if p.strip()]

    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    main_summary = analyze_company(ticker.upper())
    peer_summaries = [analyze_company(p) for p in peer_list]

    result = {
        "Target Summary": main_summary,
        "Peer Comparison": peer_summaries
    }
    return jsonify(result)

@app.route("/uploads/<path:filename>", methods=["GET"])
@require_api_key
def download_file(filename):
    return send_from_directory(app.config["UPLOAD_FOLDER"], filename, as_attachment=True)


def generate_longform_prompt(summary, peers, insights, parsed):
    prompt = f"""
### ACTIVIST REPORT: {summary['Company']} ({summary['Ticker']}) ###

1. Executive Summary
Write a 1000-word overview summarizing:
- Underperformance
- Strategic gaps
- Governance risk
- Capital efficiency
- Peer deltas
- Opportunities for shareholder value creation

2. Financial Forensics
Include full summary and peer benchmarks.

3. Capital Allocation Review
Comment on cash, debt, buybacks, CapEx, and IRR table.

4. Strategic Positioning
Pull language from parsed['strategy_flags'].

5. Operational Execution
Comment on trends in SG&A, margins, cost structure.

6. Governance & Board Review
Use parsed['board_insights'] and parsed['board_comp_table'].

7. Brand & Customer Health
Look for loyalty programs, customer metrics, channel mix.

8. Risk Heatmap
Create a list of known risks or omissions in disclosures.

9. Activist Playbook (5 Actions)
Recommend 5 bold but credible actions for shareholder value.

10. Appendix
Include IRR table, raw financials, and peer stack.

### CONTEXT DATA ###
MAIN SUMMARY:
{summary}

PEER SUMMARIES:
{peers}

INSIGHTS:
{insights}

PARSED FILE EXCERPTS:
{parsed}
"""
    return prompt

@app.route("/generate-prompt", methods=["GET"])
@require_api_key
def generate_prompt():
    ticker = request.args.get("ticker")
    peers = request.args.get("peers", "")
    peer_list = [p.strip().upper() for p in peers.split(",") if p.strip()]

    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    main_summary = analyze_company(ticker.upper())
    peer_summaries = [analyze_company(p) for p in peer_list]
    parsed = parse_uploaded_content()

    insights = []
    main_rev = main_summary["Revenue"]["value"] or 1
    for peer in peer_summaries:
        if peer["Gross Margin (%)"]["value"] and main_summary["Gross Margin (%)"]["value"] and \
           peer["Gross Margin (%)"]["value"] > main_summary["Gross Margin (%)"]["value"] + 2:
            insights.append(f"{main_summary['Ticker']} gross margin ({main_summary['Gross Margin (%)']['value']}%) is below {peer['Ticker']} at {peer['Gross Margin (%)']['value']}%.")

        if peer["SG&A as % of Revenue"]["value"] and main_summary["SG&A as % of Revenue"]["value"] and \
           peer["SG&A as % of Revenue"]["value"] < main_summary["SG&A as % of Revenue"]["value"] - 2:
            insights.append(f"{main_summary['Ticker']} SG&A % of revenue ({main_summary['SG&A as % of Revenue']['value']}%) is higher than {peer['Ticker']} at {peer['SG&A as % of Revenue']['value']}%.")

        if peer["FCF Margin (%)"]["value"] and main_summary["FCF Margin (%)"]["value"] and \
           peer["FCF Margin (%)"]["value"] > main_summary["FCF Margin (%)"]["value"] + 2:
            insights.append(f"{main_summary['Ticker']} FCF margin ({main_summary['FCF Margin (%)']['value']}%) lags {peer['Ticker']} at {peer['FCF Margin (%)']['value']}%.")

    full_prompt = generate_longform_prompt(main_summary, peer_summaries, insights, parsed)
    return jsonify({"prompt": full_prompt})

@app.route("/generate-irr", methods=["GET"])
@require_api_key
def generate_irr():
    ticker = request.args.get("ticker")
    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    data = analyze_company(ticker.upper())
    if "error" in data:
        return jsonify(data), 500

    return jsonify({"ticker": ticker.upper(), "irr_table": data.get("IRR Table", "Not Available")})

@app.route("/generate-narrative", methods=["GET"])
@require_api_key
def generate_narrative():
    ticker = request.args.get("ticker")
    peers = request.args.get("peers", "")
    peer_list = [p.strip().upper() for p in peers.split(",") if p.strip()]

    if not ticker:
        return jsonify({"error": "Missing 'ticker' parameter"}), 400

    main_summary = analyze_company(ticker.upper())
    peer_summaries = [analyze_company(p) for p in peer_list]
    parsed = parse_uploaded_content()

    insights = []
    for peer in peer_summaries:
        if peer["Gross Margin (%)"]["value"] and main_summary["Gross Margin (%)"]["value"]:
            if peer["Gross Margin (%)"]["value"] > main_summary["Gross Margin (%)"]["value"] + 2:
                insights.append(f"{main_summary['Ticker']} gross margin is below {peer['Ticker']}.")

        if peer["SG&A as % of Revenue"]["value"] and main_summary["SG&A as % of Revenue"]["value"]:
            if peer["SG&A as % of Revenue"]["value"] < main_summary["SG&A as % of Revenue"]["value"] - 2:
                insights.append(f"{main_summary['Ticker']} SG&A ratio is higher than {peer['Ticker']}.")

        if peer["FCF Margin (%)"]["value"] and main_summary["FCF Margin (%)"]["value"]:
            if peer["FCF Margin (%)"]["value"] > main_summary["FCF Margin (%)"]["value"] + 2:
                insights.append(f"{main_summary['Ticker']} FCF margin lags {peer['Ticker']}.")

    narrative = generate_longform_prompt(main_summary, peer_summaries, insights, parsed)
    return jsonify({"narrative": narrative})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(host="0.0.0.0", port=port)
